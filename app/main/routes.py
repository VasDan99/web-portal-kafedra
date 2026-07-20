from flask import render_template, redirect, url_for, flash, abort, request, send_file, session
from flask_login import login_required, current_user
from app.main import bp
from app.forms import FeedbackForm, StudentProfileForm, TeacherProfileForm, WorkUploadForm, DocumentUploadForm, \
    ChangePasswordForm, WorkMessageForm, SiteSettingsForm, AdminProfileForm
from app.models import Feedback, Student, Teacher, Discipline, Grade, Schedule, Document, WorkMessage, News, User, \
    SiteSettings
from app import db
import os
from datetime import datetime
from io import BytesIO
from docx import Document as DocxDocument
from werkzeug.utils import secure_filename
import markdown


# ==================== Публичные страницы ====================

@bp.route('/')
def index():
    breadcrumb_title = None
    news = News.query.order_by(News.created_at.desc()).limit(5).all()
    return render_template('index.html', breadcrumb_title=breadcrumb_title, news=news)


@bp.route('/about')
def about():
    breadcrumb_title = 'О факультете ИТ'
    return render_template('about.html', breadcrumb_title=breadcrumb_title)


@bp.route('/history')
def history():
    breadcrumb_title = 'История факультета'
    return render_template('history.html', breadcrumb_title=breadcrumb_title)


@bp.route('/mission')
def mission():
    breadcrumb_title = 'Миссия и цели'
    return render_template('mission.html', breadcrumb_title=breadcrumb_title)


@bp.route('/team')
def team():
    breadcrumb_title = 'Команда факультета'
    return render_template('team.html', breadcrumb_title=breadcrumb_title)


@bp.route('/departments')
def departments():
    breadcrumb_title = 'Кафедры факультета'
    return render_template('departments.html', breadcrumb_title=breadcrumb_title)


@bp.route('/teachers')
def teachers():
    breadcrumb_title = 'Преподаватели'
    teachers_list = Teacher.query.all()
    users = User.query.all()
    disciplines = Discipline.query.all()
    return render_template('teachers.html', breadcrumb_title=breadcrumb_title, teachers=teachers_list, users=users,
                           disciplines=disciplines)


@bp.route('/disciplines')
def disciplines():
    breadcrumb_title = 'Дисциплины'
    disciplines_list = Discipline.query.all()
    teachers = Teacher.query.all()
    return render_template('disciplines.html', breadcrumb_title=breadcrumb_title, disciplines=disciplines_list,
                           teachers=teachers)


@bp.route('/discipline/<int:disc_id>')
def discipline_detail(disc_id):
    discipline = Discipline.query.get_or_404(disc_id)
    breadcrumb_title = discipline.name
    teacher = Teacher.query.get(discipline.teacher_id) if discipline.teacher_id else None
    return render_template('discipline_detail.html', breadcrumb_title=breadcrumb_title, discipline=discipline,
                           teacher=teacher)


@bp.route('/schedule')
def schedule():
    breadcrumb_title = 'Расписание занятий'
    schedule_items = Schedule.query.all()
    disciplines = Discipline.query.all()
    teachers = Teacher.query.all()
    return render_template('schedule.html', breadcrumb_title=breadcrumb_title, schedule=schedule_items,
                           disciplines=disciplines, teachers=teachers)


@bp.route('/science-works')
def science_works():
    breadcrumb_title = 'Научные работы'
    return render_template('science_works.html', breadcrumb_title=breadcrumb_title)


@bp.route('/contacts')
def contacts():
    breadcrumb_title = 'Контакты'
    return render_template('contacts.html', breadcrumb_title=breadcrumb_title)


@bp.route('/feedback', methods=['GET', 'POST'])
def feedback():
    breadcrumb_title = 'Обратная связь'
    form = FeedbackForm()
    if form.validate_on_submit():
        feedback_entry = Feedback(
            name=form.name.data,
            email=form.email.data,
            subject=form.subject.data,
            message=form.message.data
        )
        db.session.add(feedback_entry)
        db.session.commit()
        flash('Ваше сообщение отправлено! Спасибо за обратную связь.', 'success')
        return redirect(url_for('main.feedback'))
    return render_template('feedback.html', breadcrumb_title=breadcrumb_title, form=form)


@bp.route('/admin')
@login_required
def admin():
    if current_user.role != 'admin':
        abort(403)
    return redirect(url_for('main.admin_dashboard'))


# ==================== Личный кабинет студента ====================

@bp.route('/cabinet/student/profile')
@login_required
def student_profile():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    if not student:
        student = Student(user_id=current_user.id, full_name=current_user.username, group_name='ИС-01', course=1)
        db.session.add(student)
        db.session.commit()
    return render_template('cabinet/student_profile.html', breadcrumb_title='Мой профиль', student=student)


@bp.route('/cabinet/student/profile/edit', methods=['GET', 'POST'])
@login_required
def student_profile_edit():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    form = StudentProfileForm()

    if form.validate_on_submit():
        student.full_name = form.full_name.data
        student.group_name = form.group_name.data
        student.course = int(form.course.data)
        student.phone = form.phone.data
        student.telegram = form.telegram.data
        student.bio = form.bio.data

        if form.avatar.data:
            avatar_file = form.avatar.data
            filename = f'user_{current_user.id}.jpg'
            os.makedirs('app/static/uploads/avatars', exist_ok=True)
            filepath = os.path.join('app/static/uploads/avatars', filename)
            avatar_file.save(filepath)
            student.avatar = f'/static/uploads/avatars/{filename}'

        db.session.commit()
        flash('Профиль обновлён!', 'success')
        return redirect(url_for('main.student_profile'))

    form.full_name.data = student.full_name
    form.group_name.data = student.group_name
    form.course.data = str(student.course)
    form.phone.data = student.phone
    form.telegram.data = student.telegram
    form.bio.data = student.bio

    return render_template('cabinet/student_profile_edit.html', breadcrumb_title='Редактирование профиля', form=form,
                           student=student)


@bp.route('/cabinet/student/grades')
@login_required
def student_grades():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    if not student:
        flash('Профиль студента не найден', 'danger')
        return redirect(url_for('main.student_profile'))

    # Получаем все оценки студента
    grades = Grade.query.filter_by(student_id=student.id).all()

    # Расчёт статистики
    average = student.get_average_grade()
    attendance_stats = student.get_attendance_stats()
    debts = student.get_debts()
    progress = student.get_progress_percentage()
    total_credits = student.get_total_credits()
    completed_credits = student.get_completed_credits()

    # Группировка оценок по дисциплинам для удобного отображения
    grades_by_discipline = {}
    for grade in grades:
        disc_name = grade.discipline.name if grade.discipline else 'Неизвестная дисциплина'
        if disc_name not in grades_by_discipline:
            grades_by_discipline[disc_name] = []
        grades_by_discipline[disc_name].append(grade)

    return render_template('cabinet/student_grades.html',
                           breadcrumb_title='Мои оценки',
                           grades=grades,
                           grades_by_discipline=grades_by_discipline,
                           student=student,
                           average=average,
                           attendance_stats=attendance_stats,
                           debts=debts,
                           progress=progress,
                           total_credits=total_credits,
                           completed_credits=completed_credits)


@bp.route('/cabinet/student/schedule')
@login_required
def student_schedule():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    schedule = Schedule.query.filter_by(group_name=student.group_name).all() if student else []
    return render_template('cabinet/student_schedule.html', breadcrumb_title='Моё расписание', schedule=schedule,
                           student=student)


@bp.route('/cabinet/student/works')
@login_required
def student_works():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    works = Document.query.filter_by(uploaded_by=current_user.id).all()
    form = WorkUploadForm()
    disciplines = Discipline.query.all()
    form.discipline_id.choices = [(d.id, d.name) for d in disciplines]
    return render_template('cabinet/student_works.html', breadcrumb_title='Мои работы',
                           student=student, form=form, works=works, disciplines=disciplines)


@bp.route('/cabinet/student/works/upload', methods=['GET', 'POST'])
@login_required
def student_works_upload():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    form = WorkUploadForm()

    disciplines = Discipline.query.all()
    form.discipline_id.choices = [(d.id, d.name) for d in disciplines]

    if form.validate_on_submit():
        file = form.file.data
        filename = f'work_{current_user.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        os.makedirs('app/static/uploads/works', exist_ok=True)
        filepath = os.path.join('app/static/uploads/works', filename)
        file.save(filepath)

        file_type = 'pdf'
        if filename.endswith('.docx'):
            file_type = 'docx'
        elif filename.endswith('.xlsx'):
            file_type = 'xlsx'

        work = Document(
            title=form.title.data,
            filename=filename,
            file_path=f'/static/uploads/works/{filename}',
            file_type=file_type,
            uploaded_by=current_user.id,
            discipline_id=form.discipline_id.data,
            is_public=False
        )
        db.session.add(work)
        db.session.commit()
        flash('Работа успешно загружена!', 'success')
        return redirect(url_for('main.student_works'))

    works = Document.query.filter_by(uploaded_by=current_user.id).all()
    return render_template('cabinet/student_works.html', breadcrumb_title='Мои работы',
                           student=student, form=form, works=works, disciplines=disciplines)


@bp.route('/cabinet/student/work/delete/<int:work_id>')
@login_required
def student_work_delete(work_id):
    if current_user.role != 'student':
        abort(403)
    work = Document.query.get_or_404(work_id)
    if work.uploaded_by == current_user.id:
        filepath = os.path.join('app', work.file_path.lstrip('/'))
        if os.path.exists(filepath):
            os.remove(filepath)
        db.session.delete(work)
        db.session.commit()
        flash('Работа удалена!', 'success')
    return redirect(url_for('main.student_works'))


@bp.route('/cabinet/student/achievements')
@login_required
def student_achievements():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    return render_template('cabinet/student_achievements.html', breadcrumb_title='Мои достижения', student=student)


@bp.route('/cabinet/student/teachers')
@login_required
def student_teachers():
    """Страница 'Мои преподаватели' для студента"""
    if current_user.role != 'student':
        abort(403)

    student = Student.query.filter_by(user_id=current_user.id).first()
    if not student:
        flash('Профиль студента не найден', 'danger')
        return redirect(url_for('main.cabinet'))

    # Получаем всех преподавателей, дисциплины и пользователей
    teachers = Teacher.query.all()
    disciplines = Discipline.query.all()
    users = User.query.all()

    return render_template('cabinet/student_teachers.html',
                           breadcrumb_title='Мои преподаватели',
                           teachers=teachers,
                           disciplines=disciplines,
                           users=users,
                           student=student)


@bp.route('/cabinet/student/ask_teacher/<int:teacher_id>', methods=['GET', 'POST'])
@login_required
def ask_teacher(teacher_id):
    if current_user.role != 'student':
        abort(403)

    teacher = Teacher.query.get_or_404(teacher_id)
    student = Student.query.filter_by(user_id=current_user.id).first()

    if request.method == 'POST':
        message = request.form.get('message')
        if message:
            feedback = Feedback(
                name=student.full_name,
                email=current_user.email,
                subject=f'Вопрос преподавателю {teacher.full_name}',
                message=message,
                status='for_teacher'
            )
            db.session.add(feedback)
            db.session.commit()
            flash('Вопрос отправлен преподавателю!', 'success')
            return redirect(url_for('main.student_teachers'))

    return render_template('cabinet/ask_teacher.html', breadcrumb_title='Задать вопрос', teacher=teacher,
                           student=student)


@bp.route('/cabinet/student/documents')
@login_required
def student_documents():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    documents = Document.query.filter_by(uploaded_by=current_user.id).all()
    form = DocumentUploadForm()
    return render_template('cabinet/student_documents.html', breadcrumb_title='Мои документы', student=student,
                           form=form, documents=documents)


@bp.route('/cabinet/student/documents/upload', methods=['GET', 'POST'])
@login_required
def student_documents_upload():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    form = DocumentUploadForm()

    if form.validate_on_submit():
        file = form.file.data
        filename = f'doc_{current_user.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{file.filename}'
        os.makedirs('app/static/uploads/documents', exist_ok=True)
        filepath = os.path.join('app/static/uploads/documents', filename)
        file.save(filepath)

        doc = Document(
            title=form.title.data,
            filename=filename,
            file_path=f'/static/uploads/documents/{filename}',
            uploaded_by=current_user.id,
            is_public=False
        )
        db.session.add(doc)
        db.session.commit()
        flash('Документ успешно загружен!', 'success')
        return redirect(url_for('main.student_documents'))

    documents = Document.query.filter_by(uploaded_by=current_user.id).all()
    return render_template('cabinet/student_documents.html', breadcrumb_title='Мои документы', student=student,
                           form=form, documents=documents)


@bp.route('/cabinet/student/document/delete/<int:doc_id>')
@login_required
def student_document_delete(doc_id):
    if current_user.role != 'student':
        abort(403)
    doc = Document.query.get_or_404(doc_id)
    if doc.uploaded_by == current_user.id:
        filepath = os.path.join('app', doc.file_path.lstrip('/'))
        if os.path.exists(filepath):
            os.remove(filepath)
        db.session.delete(doc)
        db.session.commit()
        flash('Документ удалён!', 'success')
    return redirect(url_for('main.student_documents'))


@bp.route('/cabinet/student/generate/certificate')
@login_required
def generate_certificate():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()

    doc = DocxDocument()
    doc.add_heading('Справка об обучении', 0)
    doc.add_paragraph(f'Выдана студенту: {student.full_name}')
    doc.add_paragraph(f'Группа: {student.group_name}, {student.course} курс')
    doc.add_paragraph(f'Факультет информационных технологий')
    doc.add_paragraph(f'Московский университет имени Витте')
    doc.add_paragraph(f'Дата выдачи: {datetime.now().strftime("%d.%m.%Y")}')

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return send_file(byte_io, as_attachment=True, download_name=f'spravka_{student.student_card_number}.docx')


@bp.route('/cabinet/student/settings', methods=['GET', 'POST'])
@login_required
def student_settings():
    if current_user.role != 'student':
        abort(403)
    student = Student.query.filter_by(user_id=current_user.id).first()
    password_form = ChangePasswordForm()
    profile_form = StudentProfileForm()

    if password_form.validate_on_submit() and password_form.submit.data:
        if current_user.check_password(password_form.current_password.data):
            if password_form.new_password.data == password_form.confirm_password.data:
                current_user.set_password(password_form.new_password.data)
                db.session.commit()
                flash('Пароль успешно изменён!', 'success')
            else:
                flash('Новые пароли не совпадают!', 'danger')
        else:
            flash('Неверный текущий пароль!', 'danger')
        return redirect(url_for('main.student_settings'))

    if profile_form.validate_on_submit() and profile_form.submit.data:
        student.full_name = profile_form.full_name.data
        student.group_name = profile_form.group_name.data
        student.course = int(profile_form.course.data)
        student.phone = profile_form.phone.data
        student.telegram = profile_form.telegram.data
        student.bio = profile_form.bio.data

        if profile_form.avatar.data:
            avatar_file = profile_form.avatar.data
            filename = f'avatar_{current_user.id}.jpg'
            os.makedirs('app/static/uploads/avatars', exist_ok=True)
            filepath = os.path.join('app/static/uploads/avatars', filename)
            avatar_file.save(filepath)
            student.avatar = f'/static/uploads/avatars/{filename}'

        db.session.commit()
        flash('Профиль обновлён!', 'success')
        return redirect(url_for('main.student_settings'))

    profile_form.full_name.data = student.full_name
    profile_form.group_name.data = student.group_name
    profile_form.course.data = str(student.course)
    profile_form.phone.data = student.phone
    profile_form.telegram.data = student.telegram
    profile_form.bio.data = student.bio

    return render_template('cabinet/student_settings.html', breadcrumb_title='Настройки',
                           student=student, profile_form=profile_form, password_form=password_form)


# ==================== Личный кабинет преподавателя ====================

@bp.route('/cabinet/teacher/profile')
@login_required
def teacher_profile():
    if current_user.role != 'teacher':
        abort(403)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if not teacher:
        teacher = Teacher(user_id=current_user.id, full_name=current_user.username, department='ИТ',
                          position='Преподаватель')
        db.session.add(teacher)
        db.session.commit()
    return render_template('cabinet/teacher/profile.html', breadcrumb_title='Мой профиль', teacher=teacher)


@bp.route('/cabinet/teacher/disciplines')
@login_required
def teacher_disciplines():
    if current_user.role != 'teacher':
        abort(403)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    disciplines = Discipline.query.filter_by(teacher_id=teacher.id).all() if teacher else []
    return render_template('cabinet/teacher/disciplines.html', breadcrumb_title='Мои дисциплины',
                           disciplines=disciplines, teacher=teacher)


@bp.route('/cabinet/teacher/students')
@login_required
def teacher_students():
    if current_user.role != 'teacher':
        abort(403)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    students = Student.query.all()
    return render_template('cabinet/teacher/students.html', breadcrumb_title='Студенты', students=students,
                           teacher=teacher)


@bp.route('/cabinet/teacher/students-works')
@login_required
def teacher_students_works():
    if current_user.role != 'teacher':
        abort(403)

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()

    works_query = Document.query.join(Discipline).filter(Discipline.teacher_id == teacher.id).all()

    works_data = []
    for work in works_query:
        student = Student.query.get(work.uploaded_by)
        discipline = Discipline.query.get(work.discipline_id)

        works_data.append({
            'id': work.id,
            'title': work.title,
            'file_path': work.file_path,
            'uploaded_at': work.uploaded_at,
            'student_name': student.full_name if student else 'Неизвестно',
            'discipline_name': discipline.name if discipline else '—'
        })

    return render_template('cabinet/teacher/students_works.html',
                           teacher=teacher,
                           works=works_data)


@bp.route('/cabinet/teacher/grades')
@login_required
def teacher_grades():
    if current_user.role != 'teacher':
        abort(403)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    students = Student.query.all()
    disciplines = Discipline.query.filter_by(teacher_id=teacher.id).all()
    grades = Grade.query.join(Discipline).filter(Discipline.teacher_id == teacher.id).all()
    return render_template('cabinet/teacher/grades.html', breadcrumb_title='Выставление оценок',
                           teacher=teacher, students=students, disciplines=disciplines, grades=grades)


@bp.route('/cabinet/teacher/add-grade', methods=['POST'])
@login_required
def teacher_add_grade():
    if current_user.role != 'teacher':
        abort(403)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()

    student_id = request.form.get('student_id')
    discipline_id = request.form.get('discipline_id')
    grade_value = request.form.get('grade_value')
    grade_type = request.form.get('grade_type')

    if student_id and discipline_id and grade_value:
        grade = Grade(
            student_id=student_id,
            discipline_id=discipline_id,
            grade_value=int(grade_value),
            grade_type=grade_type,
            teacher_id=teacher.id
        )
        db.session.add(grade)
        db.session.commit()
        flash('Оценка выставлена!', 'success')

    return redirect(url_for('main.teacher_grades'))


@bp.route('/cabinet/teacher/reports')
@login_required
def teacher_reports():
    if current_user.role != 'teacher':
        abort(403)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    return render_template('cabinet/teacher/reports.html', breadcrumb_title='Отчёты', teacher=teacher)


@bp.route('/cabinet/teacher/profile/edit', methods=['GET', 'POST'])
@login_required
def teacher_profile_edit():
    if current_user.role != 'teacher':
        abort(403)

    print("=== ОТЛАДКА ===")
    print("Метод:", request.method)
    print("Файлы:", request.files)

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    form = TeacherProfileForm()

    if request.method == 'POST':
        print("Это POST запрос")
        if 'avatar' in request.files:
            print("Аватар найден в request.files!")
            avatar_file = request.files['avatar']
            print(f"Имя файла: {avatar_file.filename}")
            if avatar_file.filename:
                filename = f'teacher_avatar_{current_user.id}.jpg'
                os.makedirs('app/static/uploads/avatars', exist_ok=True)
                filepath = os.path.join('app/static/uploads/avatars', filename)
                avatar_file.save(filepath)
                teacher.avatar = f'/static/uploads/avatars/{filename}'
                print(f"Файл сохранён: {filepath}")
                db.session.commit()
                flash('Аватар загружен!', 'success')
            else:
                print("Имя файла пустое!")
        else:
            print("Аватар НЕ найден в request.files!")

    if form.validate_on_submit():
        print("Форма валидна")
        teacher.full_name = form.full_name.data
        teacher.department = form.department.data
        teacher.position = form.position.data
        teacher.degree = form.degree.data
        teacher.phone = form.phone.data
        teacher.bio = form.bio.data
        db.session.commit()
        flash('Профиль обновлён!', 'success')
        return redirect(url_for('main.teacher_profile'))
    else:
        if request.method == 'POST':
            print("Форма НЕ валидна!")
            print(form.errors)

    form.full_name.data = teacher.full_name
    form.department.data = teacher.department
    form.position.data = teacher.position
    form.degree.data = teacher.degree
    form.phone.data = teacher.phone
    form.bio.data = teacher.bio

    return render_template('cabinet/teacher/profile_edit.html', breadcrumb_title='Редактирование профиля', form=form,
                           teacher=teacher)


# ==================== Чат по работам ====================

@bp.route('/cabinet/teacher/work/<int:work_id>/chat', methods=['GET', 'POST'])
@login_required
def teacher_work_chat(work_id):
    if current_user.role != 'teacher':
        abort(403)

    work = Document.query.get_or_404(work_id)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    student = Student.query.get(work.uploaded_by)

    discipline = Discipline.query.get(work.discipline_id)
    if discipline.teacher_id != teacher.id:
        abort(403)

    form = WorkMessageForm()

    if form.validate_on_submit():
        filename = None
        if form.file.data:
            file = form.file.data
            filename = f'work_{work_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{file.filename}'
            os.makedirs('app/static/uploads/reviews', exist_ok=True)
            filepath = os.path.join('app/static/uploads/reviews', filename)
            file.save(filepath)

        message = WorkMessage(
            work_id=work_id,
            from_user_id=current_user.id,
            to_user_id=student.user_id,
            message=form.message.data,
            file_path=f'/static/uploads/reviews/{filename}' if filename else None
        )
        db.session.add(message)
        db.session.commit()
        flash('Сообщение отправлено студенту!', 'success')
        return redirect(url_for('main.teacher_work_chat', work_id=work_id))

    messages = WorkMessage.query.filter_by(work_id=work_id).order_by(WorkMessage.created_at).all()

    return render_template('cabinet/teacher/work_chat.html',
                           breadcrumb_title='Обсуждение работы',
                           work=work, teacher=teacher, student=student, messages=messages, form=form)


@bp.route('/cabinet/student/work/<int:work_id>/chat', methods=['GET', 'POST'])
@login_required
def student_work_chat(work_id):
    if current_user.role != 'student':
        abort(403)

    work = Document.query.get_or_404(work_id)
    if work.uploaded_by != current_user.id:
        abort(403)

    student = Student.query.filter_by(user_id=current_user.id).first()
    discipline = Discipline.query.get(work.discipline_id)
    teacher = Teacher.query.get(discipline.teacher_id)

    form = WorkMessageForm()

    if form.validate_on_submit():
        filename = None
        if form.file.data:
            file = form.file.data
            filename = f'student_work_{work_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{file.filename}'
            os.makedirs('app/static/uploads/reviews', exist_ok=True)
            filepath = os.path.join('app/static/uploads/reviews', filename)
            file.save(filepath)

        message = WorkMessage(
            work_id=work_id,
            from_user_id=current_user.id,
            to_user_id=teacher.user_id,
            message=form.message.data,
            file_path=f'/static/uploads/reviews/{filename}' if filename else None
        )
        db.session.add(message)
        db.session.commit()
        flash('Сообщение отправлено преподавателю!', 'success')
        return redirect(url_for('main.student_work_chat', work_id=work_id))

    messages = WorkMessage.query.filter_by(work_id=work_id).order_by(WorkMessage.created_at).all()

    return render_template('cabinet/student/work_chat.html',
                           breadcrumb_title='Обсуждение работы',
                           work=work, teacher=teacher, student=student, messages=messages, form=form)


# ==================== Админ-панель ====================

@bp.route('/admin/dashboard')
@login_required
def admin_dashboard():
    if current_user.role != 'admin':
        abort(403)

    students_count = Student.query.count()
    teachers_count = Teacher.query.count()
    disciplines_count = Discipline.query.count()
    works_count = Document.query.count()

    return render_template('admin/dashboard.html',
                           breadcrumb_title='Главная',
                           students_count=students_count,
                           teachers_count=teachers_count,
                           disciplines_count=disciplines_count,
                           works_count=works_count)


@bp.route('/admin/students')
@login_required
def admin_students():
    if current_user.role != 'admin':
        abort(403)

    students = Student.query.all()
    return render_template('admin/students.html',
                           breadcrumb_title='Студенты',
                           students=students)


@bp.route('/admin/teachers')
@login_required
def admin_teachers():
    if current_user.role != 'admin':
        abort(403)

    teachers = Teacher.query.all()
    return render_template('admin/teachers.html',
                           breadcrumb_title='Преподаватели',
                           teachers=teachers)


@bp.route('/admin/schedule')
@login_required
def admin_schedule():
    if current_user.role != 'admin':
        abort(403)

    schedule = Schedule.query.all()
    disciplines = Discipline.query.all()
    teachers = Teacher.query.all()
    groups = ['ИС-01', 'ИС-02', 'ИС-03', 'ПИ-01', 'ПИ-02']
    days = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота']
    times = ['09:00-10:30', '10:45-12:15', '12:30-14:00', '14:15-15:45', '16:00-17:30']

    return render_template('admin/schedule.html',
                           breadcrumb_title='Расписание',
                           schedule=schedule,
                           disciplines=disciplines,
                           teachers=teachers,
                           groups=groups,
                           days=days,
                           times=times)


@bp.route('/admin/schedule/add', methods=['POST'])
@login_required
def admin_schedule_add():
    if current_user.role != 'admin':
        abort(403)

    discipline_id = request.form.get('discipline_id')
    group_name = request.form.get('group_name')
    day_of_week = request.form.get('day_of_week')
    lesson_time = request.form.get('lesson_time')
    classroom = request.form.get('classroom')
    teacher_id = request.form.get('teacher_id')

    days_map = {'Понедельник': 1, 'Вторник': 2, 'Среда': 3, 'Четверг': 4, 'Пятница': 5, 'Суббота': 6}
    day_num = days_map.get(day_of_week, 1)

    schedule_item = Schedule(
        discipline_id=discipline_id,
        group_name=group_name,
        day_of_week=day_num,
        lesson_time=lesson_time,
        classroom=classroom,
        teacher_id=teacher_id if teacher_id else None
    )
    db.session.add(schedule_item)
    db.session.commit()

    flash('Занятие добавлено!', 'success')
    return redirect(url_for('main.admin_schedule'))


@bp.route('/cabinet/contact_admin', methods=['GET', 'POST'])
@login_required
def contact_admin():
    form = WorkMessageForm()

    # Получаем данные пользователя в зависимости от роли
    student = None
    teacher = None

    if current_user.role == 'student':
        student = Student.query.filter_by(user_id=current_user.id).first()
    elif current_user.role == 'teacher':
        teacher = Teacher.query.filter_by(user_id=current_user.id).first()

    # Получаем историю сообщений пользователя к админу
    admin = User.query.filter_by(role='admin').first()
    messages = []
    if admin:
        messages = WorkMessage.query.filter_by(
            from_user_id=current_user.id,
            to_user_id=admin.id
        ).order_by(WorkMessage.created_at).all()

    if form.validate_on_submit():
        filename = None
        if form.file.data:
            file = form.file.data
            filename = f'msg_{current_user.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{file.filename}'
            os.makedirs('app/static/uploads/messages', exist_ok=True)
            filepath = os.path.join('app/static/uploads/messages', filename)
            file.save(filepath)

        message = WorkMessage(
            work_id=None,
            from_user_id=current_user.id,
            to_user_id=admin.id if admin else None,
            message=form.message.data,
            file_path=f'/static/uploads/messages/{filename}' if filename else None
        )
        db.session.add(message)
        db.session.commit()
        flash('Ваше сообщение отправлено администратору!', 'success')
        return redirect(url_for('main.contact_admin'))

    return render_template('cabinet/contact_admin.html',
                           breadcrumb_title='Связь с администратором',
                           form=form,
                           student=student,
                           teacher=teacher,
                           messages=messages)


@bp.route('/admin/schedule/delete/<int:schedule_id>')
@login_required
def admin_schedule_delete(schedule_id):
    if current_user.role != 'admin':
        abort(403)

    schedule_item = Schedule.query.get_or_404(schedule_id)
    db.session.delete(schedule_item)
    db.session.commit()

    flash('Занятие удалено!', 'success')
    return redirect(url_for('main.admin_schedule'))


@bp.route('/admin/news')
@login_required
def admin_news():
    if current_user.role != 'admin':
        abort(403)

    news = News.query.order_by(News.created_at.desc()).all()
    return render_template('admin/news.html',
                           breadcrumb_title='Новости',
                           news=news)


@bp.route('/admin/feedback')
@login_required
def admin_feedback():
    if current_user.role != 'admin':
        abort(403)

    messages = Feedback.query.order_by(Feedback.created_at.desc()).all()
    return render_template('admin/feedback.html',
                           breadcrumb_title='Сообщения',
                           messages=messages)


@bp.route('/admin/student/add', methods=['GET', 'POST'])
@login_required
def admin_student_add():
    if current_user.role != 'admin':
        abort(403)

    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        full_name = request.form.get('full_name')
        group_name = request.form.get('group_name')
        course = request.form.get('course')
        phone = request.form.get('phone')

        if User.query.filter_by(username=username).first():
            flash('Пользователь с таким логином уже существует!', 'danger')
            return redirect(url_for('main.admin_student_add'))

        user = User(username=username, email=email, role='student')
        user.set_password(password)
        db.session.add(user)
        db.session.flush()

        student = Student(
            user_id=user.id,
            full_name=full_name,
            group_name=group_name,
            course=int(course),
            phone=phone
        )
        db.session.add(student)
        db.session.commit()

        flash('Студент успешно добавлен!', 'success')
        return redirect(url_for('main.admin_students'))

    return render_template('admin/student_add.html', breadcrumb_title='Добавление студента')


@bp.route('/admin/student/edit/<int:student_id>', methods=['GET', 'POST'])
@login_required
def admin_student_edit(student_id):
    if current_user.role != 'admin':
        abort(403)

    student = Student.query.get_or_404(student_id)

    if request.method == 'POST':
        student.full_name = request.form.get('full_name')
        student.group_name = request.form.get('group_name')
        student.course = int(request.form.get('course'))
        student.phone = request.form.get('phone')

        if request.form.get('password'):
            student.user.set_password(request.form.get('password'))

        db.session.commit()
        flash('Студент обновлён!', 'success')
        return redirect(url_for('main.admin_students'))

    return render_template('admin/student_edit.html', breadcrumb_title='Редактирование студента', student=student)


@bp.route('/admin/student/delete/<int:student_id>')
@login_required
def admin_student_delete(student_id):
    if current_user.role != 'admin':
        abort(403)

    student = Student.query.get_or_404(student_id)
    user = student.user
    db.session.delete(student)
    db.session.delete(user)
    db.session.commit()

    flash('Студент удалён!', 'success')
    return redirect(url_for('main.admin_students'))


@bp.route('/admin/teacher/add', methods=['GET', 'POST'])
@login_required
def admin_teacher_add():
    if current_user.role != 'admin':
        abort(403)

    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        full_name = request.form.get('full_name')
        department = request.form.get('department')
        position = request.form.get('position')
        degree = request.form.get('degree')
        phone = request.form.get('phone')

        if User.query.filter_by(username=username).first():
            flash('Пользователь с таким логином уже существует!', 'danger')
            return redirect(url_for('main.admin_teacher_add'))

        user = User(username=username, email=email, role='teacher')
        user.set_password(password)
        db.session.add(user)
        db.session.flush()

        teacher = Teacher(
            user_id=user.id,
            full_name=full_name,
            department=department,
            position=position,
            degree=degree,
            phone=phone
        )
        db.session.add(teacher)
        db.session.commit()

        flash('Преподаватель успешно добавлен!', 'success')
        return redirect(url_for('main.admin_teachers'))

    return render_template('admin/teacher_add.html', breadcrumb_title='Добавление преподавателя')


@bp.route('/admin/teacher/edit/<int:teacher_id>', methods=['GET', 'POST'])
@login_required
def admin_teacher_edit(teacher_id):
    if current_user.role != 'admin':
        abort(403)

    teacher = Teacher.query.get_or_404(teacher_id)

    if request.method == 'POST':
        teacher.full_name = request.form.get('full_name')
        teacher.department = request.form.get('department')
        teacher.position = request.form.get('position')
        teacher.degree = request.form.get('degree')
        teacher.phone = request.form.get('phone')

        if request.form.get('password'):
            teacher.user.set_password(request.form.get('password'))

        db.session.commit()
        flash('Преподаватель обновлён!', 'success')
        return redirect(url_for('main.admin_teachers'))

    return render_template('admin/teacher_edit.html', breadcrumb_title='Редактирование преподавателя', teacher=teacher)


@bp.route('/admin/teacher/delete/<int:teacher_id>')
@login_required
def admin_teacher_delete(teacher_id):
    if current_user.role != 'admin':
        abort(403)

    teacher = Teacher.query.get_or_404(teacher_id)
    user = teacher.user
    db.session.delete(teacher)
    db.session.delete(user)
    db.session.commit()

    flash('Преподаватель удалён!', 'success')
    return redirect(url_for('main.admin_teachers'))


@bp.route('/admin/news/add', methods=['GET', 'POST'])
@login_required
def admin_news_add():
    if current_user.role != 'admin':
        abort(403)

    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')

        image_url = None
        if 'image' in request.files:
            file = request.files['image']
            if file and file.filename:
                filename = secure_filename(f'news_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{file.filename}')
                os.makedirs('app/static/uploads/news', exist_ok=True)
                filepath = os.path.join('app/static/uploads/news', filename)
                file.save(filepath)
                image_url = f'/static/uploads/news/{filename}'

        news = News(title=title, content=content, image_url=image_url, author_id=current_user.id)
        db.session.add(news)
        db.session.commit()

        flash('Новость добавлена!', 'success')
        return redirect(url_for('main.admin_news'))

    return render_template('admin/news_add.html', breadcrumb_title='Добавление новости')


@bp.route('/admin/news/edit/<int:news_id>', methods=['GET', 'POST'])
@login_required
def admin_news_edit(news_id):
    if current_user.role != 'admin':
        abort(403)

    news = News.query.get_or_404(news_id)

    if request.method == 'POST':
        news.title = request.form.get('title')
        news.content = request.form.get('content')

        if 'image' in request.files:
            file = request.files['image']
            if file and file.filename:
                filename = secure_filename(f'news_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{file.filename}')
                os.makedirs('app/static/uploads/news', exist_ok=True)
                filepath = os.path.join('app/static/uploads/news', filename)
                file.save(filepath)
                news.image_url = f'/static/uploads/news/{filename}'

        db.session.commit()
        flash('Новость обновлена!', 'success')
        return redirect(url_for('main.admin_news'))

    return render_template('admin/news_edit.html', breadcrumb_title='Редактирование новости', news=news)


@bp.route('/admin/news/delete/<int:news_id>')
@login_required
def admin_news_delete(news_id):
    if current_user.role != 'admin':
        abort(403)

    news = News.query.get_or_404(news_id)
    db.session.delete(news)
    db.session.commit()

    flash('Новость удалена!', 'success')
    return redirect(url_for('main.admin_news'))


@bp.route('/admin/feedback/reply/<int:feedback_id>', methods=['GET', 'POST'])
@login_required
def admin_feedback_reply(feedback_id):
    if current_user.role != 'admin':
        abort(403)

    feedback = Feedback.query.get_or_404(feedback_id)

    if request.method == 'POST':
        reply = request.form.get('reply')
        if reply:
            feedback.reply = reply
            feedback.replied_at = datetime.utcnow()
            feedback.status = 'replied'
            db.session.commit()
            flash('Ответ отправлен!', 'success')

            # Здесь можно добавить отправку email пользователю
            # send_mail(feedback.email, f'Ответ на ваше обращение: {feedback.subject}', reply)

            return redirect(url_for('main.admin_feedback'))

    return render_template('admin/feedback_reply.html', breadcrumb_title='Ответ на сообщение', feedback=feedback)


@bp.route('/news/<int:news_id>')
def news_detail(news_id):
    news_item = News.query.get_or_404(news_id)
    breadcrumb_title = news_item.title
    # Преобразуем Markdown в HTML
    news_item.content_html = markdown.markdown(news_item.content, extensions=['extra'])
    return render_template('news_detail.html', breadcrumb_title=breadcrumb_title, news=news_item)


@bp.route('/admin/all-messages')
@login_required
def admin_all_messages():
    if current_user.role != 'admin':
        abort(403)

    # Сообщения из обратной связи (от гостей)
    feedback_messages = Feedback.query.order_by(Feedback.created_at.desc()).all()

    # Личные сообщения пользователей (студентов и преподавателей)
    personal_messages = WorkMessage.query.filter(
        WorkMessage.to_user_id == current_user.id,
        WorkMessage.work_id.is_(None)
    ).order_by(WorkMessage.created_at.desc()).all()

    return render_template('admin/all_messages.html',
                           breadcrumb_title='Все сообщения',
                           feedback_messages=feedback_messages,
                           personal_messages=personal_messages)


@bp.route('/admin/personal-message/reply/<int:message_id>', methods=['POST'])
@login_required
def admin_personal_message_reply(message_id):
    if current_user.role != 'admin':
        abort(403)

    msg = WorkMessage.query.get_or_404(message_id)
    reply_text = request.form.get('reply')

    if reply_text:
        msg.reply = reply_text
        msg.replied_at = datetime.utcnow()
        msg.is_read = True
        db.session.commit()
        flash('Ответ отправлен пользователю!', 'success')

    return redirect(url_for('main.admin_all_messages'))


@bp.route('/admin/feedback/delete/<int:feedback_id>')
@login_required
def admin_feedback_delete(feedback_id):
    if current_user.role != 'admin':
        abort(403)

    feedback = Feedback.query.get_or_404(feedback_id)
    db.session.delete(feedback)
    db.session.commit()

    flash('Сообщение удалено!', 'success')
    return redirect(url_for('main.admin_all_messages'))


@bp.route('/admin/personal-message/delete/<int:message_id>')
@login_required
def admin_personal_message_delete(message_id):
    if current_user.role != 'admin':
        abort(403)

    message = WorkMessage.query.get_or_404(message_id)
    db.session.delete(message)
    db.session.commit()

    flash('Сообщение удалено!', 'success')
    return redirect(url_for('main.admin_all_messages'))


@bp.route('/language/<lang>')
def set_language(lang):
    if lang in ['ru', 'en', 'zh']:
        session['lang'] = lang
    return redirect(request.referrer or url_for('main.index'))


# ==================== Генерация отчётов для преподавателя ====================

@bp.route('/cabinet/teacher/generate/grades-docx')
@login_required
def teacher_generate_grades_docx():
    if current_user.role != 'teacher':
        abort(403)

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    disciplines = Discipline.query.filter_by(teacher_id=teacher.id).all()

    doc = DocxDocument()
    doc.add_heading(f'Ведомость успеваемости преподавателя {teacher.full_name}', 0)
    doc.add_paragraph(f'Кафедра: {teacher.department}')
    doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')

    for discipline in disciplines:
        doc.add_heading(f'Дисциплина: {discipline.name}', level=1)
        doc.add_paragraph(f'Код: {discipline.code}, Часов: {discipline.hours}')

        # Получаем оценки по дисциплине
        grades = Grade.query.filter_by(discipline_id=discipline.id).all()

        if grades:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Студент'
            hdr_cells[1].text = 'Группа'
            hdr_cells[2].text = 'Оценка'
            hdr_cells[3].text = 'Дата'

            for grade in grades:
                student = Student.query.get(grade.student_id)
                if student:
                    row_cells = table.add_row().cells
                    row_cells[0].text = student.full_name
                    row_cells[1].text = student.group_name
                    row_cells[2].text = str(grade.grade_value)
                    row_cells[3].text = grade.date.strftime('%d.%m.%Y') if grade.date else '—'
        else:
            doc.add_paragraph('Оценок пока нет.')

        doc.add_paragraph()

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return send_file(byte_io, as_attachment=True,
                     download_name=f'vedomost_{teacher.full_name}_{datetime.now().strftime("%Y%m%d")}.docx')


@bp.route('/cabinet/teacher/generate/students-xlsx')
@login_required
def teacher_generate_students_xlsx():
    if current_user.role != 'teacher':
        abort(403)

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    students = Student.query.all()

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Студенты"

    # Заголовки
    headers = ['ID', 'ФИО', 'Группа', 'Курс', 'Email', 'Телефон']
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')

    # Данные
    for row, student in enumerate(students, 2):
        ws.cell(row=row, column=1, value=student.id)
        ws.cell(row=row, column=2, value=student.full_name)
        ws.cell(row=row, column=3, value=student.group_name)
        ws.cell(row=row, column=4, value=student.course)
        ws.cell(row=row, column=5, value=student.user.email if student.user else '—')
        ws.cell(row=row, column=6, value=student.phone or '—')

    # Автоширина колонок
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 30)
        ws.column_dimensions[col_letter].width = adjusted_width

    byte_io = BytesIO()
    wb.save(byte_io)
    byte_io.seek(0)

    return send_file(byte_io, as_attachment=True, download_name=f'students_{datetime.now().strftime("%Y%m%d")}.xlsx')


@bp.route('/cabinet/teacher/generate/grades-xlsx')
@login_required
def teacher_generate_grades_xlsx():
    if current_user.role != 'teacher':
        abort(403)

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    grades = Grade.query.join(Discipline).filter(Discipline.teacher_id == teacher.id).all()

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Оценки"

    # Заголовки
    headers = ['ID', 'Студент', 'Группа', 'Дисциплина', 'Оценка', 'Тип', 'Дата']
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')

    # Данные
    for row, grade in enumerate(grades, 2):
        student = Student.query.get(grade.student_id)
        discipline = Discipline.query.get(grade.discipline_id)

        ws.cell(row=row, column=1, value=grade.id)
        ws.cell(row=row, column=2, value=student.full_name if student else '—')
        ws.cell(row=row, column=3, value=student.group_name if student else '—')
        ws.cell(row=row, column=4, value=discipline.name if discipline else '—')
        ws.cell(row=row, column=5, value=grade.grade_value)
        ws.cell(row=row, column=6, value=grade.grade_type or '—')
        ws.cell(row=row, column=7, value=grade.date.strftime('%d.%m.%Y') if grade.date else '—')

    # Автоширина колонок
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 30)
        ws.column_dimensions[col_letter].width = adjusted_width

    byte_io = BytesIO()
    wb.save(byte_io)
    byte_io.seek(0)

    return send_file(byte_io, as_attachment=True, download_name=f'grades_{datetime.now().strftime("%Y%m%d")}.xlsx')


@bp.route('/admin/settings', methods=['GET', 'POST'])
@login_required
def admin_settings():
    if current_user.role != 'admin':
        abort(403)

    settings = SiteSettings.query.first()
    if not settings:
        settings = SiteSettings()
        db.session.add(settings)
        db.session.commit()

    form = SiteSettingsForm()

    if form.validate_on_submit():
        settings.site_title = form.site_title.data
        settings.site_description = form.site_description.data
        settings.email = form.email.data
        settings.phone = form.phone.data
        settings.address = form.address.data
        settings.work_hours = form.work_hours.data
        settings.vk_url = form.vk_url.data
        settings.telegram_url = form.telegram_url.data
        settings.max_url = form.max_url.data
        settings.primary_color = form.primary_color.data
        settings.secondary_color = form.secondary_color.data
        settings.accent_color = form.accent_color.data
        settings.about_text = form.about_text.data

        # Обработка загрузки логотипа
        logo_file = request.files.get('logo')
        if logo_file and logo_file.filename:
            # Проверяем расширение файла
            allowed_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.svg'}
            file_ext = os.path.splitext(logo_file.filename)[1].lower()
            if file_ext in allowed_extensions:
                filename = f'logo_{datetime.now().strftime("%Y%m%d_%H%M%S")}{file_ext}'
                upload_folder = os.path.join('app/static/uploads')
                os.makedirs(upload_folder, exist_ok=True)
                filepath = os.path.join(upload_folder, filename)
                logo_file.save(filepath)
                settings.logo_path = f'/static/uploads/{filename}'
                flash('Логотип успешно загружен!', 'success')
            else:
                flash('Неподдерживаемый формат файла. Используйте PNG, JPG, GIF или SVG.', 'danger')

        db.session.commit()
        flash('Настройки сохранены!', 'success')
        return redirect(url_for('main.admin_settings'))

    # Заполняем форму текущими данными
    form.site_title.data = settings.site_title
    form.site_description.data = settings.site_description
    form.email.data = settings.email
    form.phone.data = settings.phone
    form.address.data = settings.address
    form.work_hours.data = settings.work_hours
    form.vk_url.data = settings.vk_url
    form.telegram_url.data = settings.telegram_url
    form.max_url.data = settings.max_url
    form.primary_color.data = settings.primary_color
    form.secondary_color.data = settings.secondary_color
    form.accent_color.data = settings.accent_color
    form.about_text.data = settings.about_text

    return render_template('admin/settings.html', breadcrumb_title='Настройки сайта', form=form, settings=settings)


@bp.route('/admin/profile', methods=['GET', 'POST'])
@login_required
def admin_profile():
    if current_user.role != 'admin':
        abort(403)

    form = AdminProfileForm()

    if form.validate_on_submit():
        if form.username.data:
            current_user.username = form.username.data
        if form.email.data:
            current_user.email = form.email.data
        if form.password.data:
            current_user.set_password(form.password.data)

        db.session.commit()
        flash('Профиль обновлён!', 'success')
        return redirect(url_for('main.admin_profile'))

    form.username.data = current_user.username
    form.email.data = current_user.email

    return render_template('admin/profile.html', breadcrumb_title='Профиль администратора', form=form)